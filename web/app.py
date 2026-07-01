"""
Flask web UI for Denial Defense multi-agent harness.
"""

from __future__ import annotations

import json
import os
import queue
import sys
import threading
import time
import uuid
from concurrent.futures import ThreadPoolExecutor
from functools import wraps
from io import BytesIO
from pathlib import Path

from dotenv import load_dotenv
from flask import Flask, Response, jsonify, render_template, request, send_file

load_dotenv()

sys.path.insert(0, str(Path(__file__).parent.parent))

from agents.baseline import single_agent_appeal
from agents.config import get_settings
from agents.harness import run_harness
from agents.logging_config import log
from agents.progress import register_callback, unregister_callback

settings = get_settings()
app = Flask(__name__, template_folder="templates")

DEMO_DIR = Path(__file__).parent.parent / "data" / "demo"
CASES: dict = {}
for case_file in DEMO_DIR.glob("case_*.json"):
    try:
        CASES[case_file.stem] = json.loads(case_file.read_text(encoding="utf-8"))
    except Exception as e:
        log.warning("Could not load %s: %s", case_file, e)

# In-memory SSE progress queues keyed by job_id
_progress_queues: dict[str, queue.Queue] = {}


def _require_api_key(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if settings.flask_api_key:
            provided = request.headers.get("X-API-Key") or request.args.get("api_key")
            if provided != settings.flask_api_key:
                return jsonify({"status": "error", "message": "Unauthorized"}), 401
        return f(*args, **kwargs)

    return decorated


def _enrich_metrics(metrics: dict, case: dict, denial_text: str) -> dict:
    metrics = dict(metrics or {})
    pre_labeled = case.get("denial_letter", {}).get("inferred_carc", [])
    if not metrics.get("carc_codes_matched") and pre_labeled:
        metrics["carc_codes_matched"] = [c.replace("CARC_", "") for c in pre_labeled]
    if not metrics.get("federal_protections_found"):
        denial_lower = denial_text.lower()
        protections = []
        if any(kw in denial_lower for kw in ["parity", "mental health", "substance use", "sud", "mhpaea"]):
            protections.append("mental_health_parity (MHPAEA)")
        if any(kw in denial_lower for kw in ["out-of-network", "out of network", "emergency"]):
            protections.append("no_surprises_act")
        if any(kw in denial_lower for kw in ["gender", "transgender"]):
            protections.append("ACA_section_1557")
        metrics["federal_protections_found"] = protections
    return metrics


def _format_harness_response(case: dict, result: dict) -> dict:
    denial_text = case["denial_letter"]["denial_text"]
    metrics = _enrich_metrics(result.get("_metrics", {}), case, denial_text)
    return {
        "status": "ok",
        "case_name": case.get("display_name", "Custom Case"),
        "denial": case["denial_letter"],
        "patient_context": case["patient_context"],
        "round_1": result.get("round_1") or result.get("round_1_outputs", {}),
        "round_2": result.get("round_2") or result.get("round_2_outputs", {}),
        "critiques": result.get("critiques", []),
        "final_verdict": result.get("final_verdict", ""),
        "metrics": metrics,
        "evidence_checklist": result.get("_evidence_checklist", []),
        "imr_cases": result.get("_imr_cases", []),
    }


def _error_response(message: str, status: int = 500):
    log.error(message)
    return jsonify({"status": "error", "message": message}), status


@app.route("/")
def index():
    return render_template("index.html", cases=CASES)


@app.route("/cases")
def list_cases():
    return jsonify({cid: c["display_name"] for cid, c in CASES.items()})


@app.route("/run/<case_id>", methods=["POST"])
@_require_api_key
def run_case(case_id):
    if case_id not in CASES:
        return _error_response(f"Case '{case_id}' not found", 404)
    case = CASES[case_id]
    try:
        result = run_harness(
            denial_letter=case["denial_letter"]["denial_text"],
            patient_context=json.dumps(case["patient_context"], indent=2),
        )
        return jsonify(_format_harness_response(case, result))
    except Exception as e:
        return _error_response(f"Harness failed: {e}")


@app.route("/run/custom", methods=["POST"])
@_require_api_key
def run_custom():
    data = request.get_json(silent=True) or {}
    denial_text = (data.get("denial_letter") or "").strip()
    patient_ctx = data.get("patient_context")
    if not denial_text:
        return _error_response("denial_letter is required", 400)
    if isinstance(patient_ctx, dict):
        patient_ctx_str = json.dumps(patient_ctx, indent=2)
    else:
        patient_ctx_str = str(patient_ctx or "").strip()
    if not patient_ctx_str:
        return _error_response("patient_context is required", 400)

    case = {
        "display_name": data.get("display_name", "Custom Case"),
        "denial_letter": {"denial_text": denial_text, "insurer": data.get("insurer", "Unknown")},
        "patient_context": patient_ctx if isinstance(patient_ctx, dict) else {"notes": patient_ctx_str},
    }
    try:
        result = run_harness(denial_letter=denial_text, patient_context=patient_ctx_str)
        return jsonify(_format_harness_response(case, result))
    except Exception as e:
        return _error_response(f"Harness failed: {e}")


@app.route("/compare/<case_id>", methods=["POST"])
@_require_api_key
def compare_case(case_id):
    if case_id not in CASES:
        return _error_response(f"Case '{case_id}' not found", 404)

    case = CASES[case_id]
    denial_text = case["denial_letter"]["denial_text"]
    patient_ctx = json.dumps(case["patient_context"], indent=2)

    try:
        with ThreadPoolExecutor(max_workers=2) as pool:
            baseline_future = pool.submit(single_agent_appeal, denial_text, patient_ctx)
            harness_future = pool.submit(run_harness, denial_text, patient_ctx)
            baseline_output = baseline_future.result()
            harness_output = harness_future.result()

        metrics = _enrich_metrics(harness_output.get("_metrics", {}), case, denial_text)
        formatted = _format_harness_response(case, harness_output)

        return jsonify(
            {
                "status": "ok",
                "denial": case["denial_letter"],
                "case_display_name": case.get("display_name", case_id),
                "baseline": {
                    "appeal_letter": baseline_output,
                    "agents_used": 1,
                    "rounds": 0,
                    "label": "Single-Agent Baseline (one GPT-4o call, one draft)",
                },
                "harness": {
                    **formatted,
                    "agents_used": 5,
                    "rounds": metrics.get("rounds_completed", 2),
                    "label": "Multi-Agent Harness with Adversarial Critic",
                    "metrics": metrics,
                },
            }
        )
    except Exception as e:
        return _error_response(f"Comparison failed: {e}")


@app.route("/run/<case_id>/stream", methods=["POST"])
@_require_api_key
def run_case_stream(case_id):
    if case_id not in CASES:
        return _error_response(f"Case '{case_id}' not found", 404)

    job_id = str(uuid.uuid4())
    q: queue.Queue = queue.Queue()
    _progress_queues[job_id] = q
    case = CASES[case_id]

    def progress_cb(step: str, detail: str = ""):
        q.put({"step": step, "detail": detail, "ts": time.time()})

    def worker():
        register_callback(progress_cb)
        try:
            result = run_harness(
                denial_letter=case["denial_letter"]["denial_text"],
                patient_context=json.dumps(case["patient_context"], indent=2),
            )
            q.put({"done": True, "result": _format_harness_response(case, result)})
        except Exception as e:
            q.put({"done": True, "error": str(e)})
        finally:
            unregister_callback(progress_cb)

    threading.Thread(target=worker, daemon=True).start()

    def generate():
        yield f"data: {json.dumps({'job_id': job_id, 'step': 'queued'})}\n\n"
        while True:
            try:
                item = q.get(timeout=120)
            except queue.Empty:
                yield f"data: {json.dumps({'step': 'heartbeat'})}\n\n"
                continue
            yield f"data: {json.dumps(item)}\n\n"
            if item.get("done"):
                break
        _progress_queues.pop(job_id, None)

    return Response(generate(), mimetype="text/event-stream")


@app.route("/export", methods=["POST"])
@_require_api_key
def export_appeal():
    data = request.get_json(silent=True) or {}
    text = (data.get("appeal_text") or "").strip()
    fmt = (data.get("format") or "txt").lower()
    filename = data.get("filename", "appeal_letter")

    if not text:
        return _error_response("appeal_text is required", 400)

    if fmt == "docx":
        try:
            from docx import Document

            doc = Document()
            doc.add_heading("Insurance Appeal Letter", level=1)
            for para in text.split("\n\n"):
                doc.add_paragraph(para)
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return send_file(
                buf,
                as_attachment=True,
                download_name=f"{filename}.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except ImportError:
            return _error_response("python-docx not installed", 501)

    buf = BytesIO(text.encode("utf-8"))
    return send_file(buf, as_attachment=True, download_name=f"{filename}.txt", mimetype="text/plain")


if __name__ == "__main__":
    if not settings.openai_api_key:
        print("ERROR: OPENAI_API_KEY required")
        sys.exit(1)

    print("=" * 80)
    print("Denial Defense - Multi-Agent Insurance Appeal Harness")
    print(f"Cases loaded: {len(CASES)}")
    print(f"http://{settings.flask_host}:{settings.flask_port}")
    print("=" * 80)
    app.run(host=settings.flask_host, port=settings.flask_port, debug=settings.flask_debug)
